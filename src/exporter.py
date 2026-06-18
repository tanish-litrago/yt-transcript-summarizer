"""
src/exporter.py — v1.1 Complete Rewrite
Fixes: Generated date cutoff, white pages, page breaks mid-content,
       dark background on all pages, proper meta line rendering.
"""

import os, sys, re
from datetime import datetime

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from config import OUTPUT_DIR


def _safe_filename(video_id):
    return f"notes_{video_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"


def _clean(text):
    replacements = {
        "\u2014": "--",  "\u2013": "-",   "\u2012": "-",  "\u2015": "--",
        "\u2010": "-",   "\u2011": "-",   "\u2018": "'",  "\u2019": "'",
        "\u201A": "'",   "\u201B": "'",   "\u201C": '"',  "\u201D": '"',
        "\u201E": '"',   "\u201F": '"',   "\u2026": "...","\u2022": "-",
        "\u00B7": "-",   "\u00A0": " ",   "\u00AD": "-",  "\u2039": "<",
        "\u203A": ">",   "\u2032": "'",   "\u2044": "/",  "\u2060": "",
        "\u200B": "",    "\u200C": "",    "\u200D": "",   "\uFEFF": "",
        "\U0001F511": "", "\U0001F464": "", "\U0001F4DD": "",
        "\U0001F4C4": "", "\U0001F4D6": "",
    }
    for char, rep in replacements.items():
        text = text.replace(char, rep)
    result = ""
    for ch in text:
        try:
            ch.encode("latin-1")
            result += ch
        except UnicodeEncodeError:
            result += ""
    return result


def _strip_md(text):
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"`(.+?)`",        r"\1", text)
    return text


def _c(text):
    return _clean(_strip_md(text))


def export_markdown(notes_md, video_id):
    path = os.path.join(OUTPUT_DIR, _safe_filename(video_id) + ".md")
    with open(path, "w", encoding="utf-8") as f:
        f.write(notes_md)
    print(f"[Export] Markdown saved: {path}")
    return path


def export_pdf(notes_md, video_id):
    try:
        from fpdf import FPDF
    except ImportError:
        raise RuntimeError("fpdf2 not installed.")

    BG       = (11,  15,  25)
    BG2      = (19,  27,  46)
    ACCENT   = (99,  102, 241)
    ACCENT_L = (129, 140, 248)
    BODY     = (226, 232, 240)
    BODY2    = (148, 163, 184)
    GREEN    = (16,  185, 129)
    BORDER   = (30,  41,  59)

    class PDF(FPDF):
        def header(self):
            self.set_fill_color(*BG)
            self.rect(0, 0, 210, 297, "F")

        def footer(self):
            self.set_y(-12)
            self.set_font("Helvetica", "I", 7)
            self.set_text_color(*BODY2)
            self.cell(0, 8, f"Page {self.page_no()} | YouTube Transcript Summarizer", align="C")

    pdf = PDF(orientation="P", unit="mm", format="A4")
    pdf.set_margins(left=18, top=22, right=18)
    pdf.set_auto_page_break(auto=True, margin=18)
    pdf.add_page()

    W = 174

    def h1(text):
        text = _c(text)
        if not text.strip(): return
        pdf.set_font("Helvetica", "B", 17)
        pdf.set_text_color(*ACCENT)
        pdf.multi_cell(W, 9, text, align="L")
        y = pdf.get_y() + 1
        pdf.set_draw_color(*ACCENT)
        pdf.set_line_width(0.4)
        pdf.line(18, y, 192, y)
        pdf.ln(4)
        pdf.set_text_color(*BODY)

    def h2(text):
        text = _c(text)
        if not text.strip(): return
        pdf.ln(3)
        y = pdf.get_y()
        pdf.set_fill_color(*BG2)
        pdf.rect(16, y - 1, W + 2, 8, "F")
        pdf.set_font("Helvetica", "B", 11)
        pdf.set_text_color(*ACCENT)
        pdf.set_x(20)
        pdf.cell(W, 7, text, align="L")
        pdf.ln(8)
        pdf.set_text_color(*BODY)

    def h3(text):
        text = _c(text)
        if not text.strip(): return
        pdf.ln(2)
        pdf.set_font("Helvetica", "B", 10)
        pdf.set_text_color(*ACCENT_L)
        pdf.multi_cell(W, 6, text, align="L")
        pdf.set_text_color(*BODY)

    def body(text):
        text = _c(text)
        if not text.strip(): return
        pdf.set_font("Helvetica", "", 9)
        pdf.set_text_color(*BODY)
        pdf.multi_cell(W, 5.5, text, align="L")

    def meta(key, val):
        key = _clean(key)
        val = _clean(val)
        if not val.strip(): return
        pdf.set_font("Helvetica", "B", 9)
        pdf.set_text_color(*ACCENT_L)
        kw = pdf.get_string_width(key + ":  ")
        pdf.cell(kw, 6, key + ": ")
        pdf.set_font("Helvetica", "", 9)
        pdf.set_text_color(*BODY2)
        pdf.multi_cell(W - kw, 6, val, align="L")

    def bullet(text):
        text = _c(text)
        if not text.strip(): return
        y = pdf.get_y()
        pdf.set_fill_color(*GREEN)
        pdf.ellipse(20.5, y + 1.8, 1.8, 1.8, "F")
        pdf.set_font("Helvetica", "", 9)
        pdf.set_text_color(*BODY)
        pdf.set_x(24)
        pdf.multi_cell(W - 6, 5.5, text, align="L")

    def keyword_tags(text):
        text = _c(text)
        tags = [t.strip() for t in re.split(r"\s{2,}", text) if t.strip()]
        if not tags:
            tags = [t for t in text.split() if t.strip()]
        if not tags: return
        x0 = 18
        y  = pdf.get_y()
        x  = x0
        pdf.set_font("Helvetica", "", 8)
        for tag in tags:
            tw = pdf.get_string_width(tag) + 8
            if x + tw > 190:
                x = x0
                y += 7
            if y > 272:
                pdf.add_page()
                y = pdf.get_y()
                x = x0
            pdf.set_xy(x, y)
            pdf.set_fill_color(30, 27, 75)
            pdf.set_draw_color(*ACCENT)
            pdf.set_line_width(0.2)
            pdf.rect(x, y, tw, 6, "FD")
            pdf.set_text_color(*ACCENT_L)
            pdf.set_xy(x + 4, y + 0.5)
            pdf.cell(tw - 8, 5, tag)
            x += tw + 4
        pdf.set_y(y + 8)
        pdf.set_text_color(*BODY)

    def hr():
        pdf.ln(3)
        pdf.set_draw_color(*BORDER)
        pdf.set_line_width(0.25)
        y = pdf.get_y()
        pdf.line(18, y, 192, y)
        pdf.ln(4)

    in_keywords = False

    for raw in notes_md.splitlines():
        line = raw.rstrip()
        try:
            if line.startswith("# "):
                in_keywords = False
                h1(line[2:])

            elif line.startswith("## "):
                heading = _c(line[3:])
                in_keywords = any(x in heading.lower() for x in
                                  ["key terms", "named entities", "keywords"])
                h2(heading)

            elif line.startswith("### "):
                in_keywords = False
                h3(line[4:])

            elif line.startswith("- "):
                content = re.sub(r"\*\*(.+?)\*\*", r"\1", line[2:])
                bullet(content)

            elif line.startswith("**") and (":**" in line or ": " in line):
                match = re.match(r"\*\*(.+?)\*\*:?\s*(.*)", line)
                if match:
                    meta(match.group(1), match.group(2))
                else:
                    body(line)

            elif line == "---":
                hr()

            elif line == "":
                pdf.ln(2)

            else:
                if "`" in line:
                    cleaned = re.sub(r"`(.+?)`", r"\1", line)
                    if in_keywords:
                        keyword_tags(cleaned)
                    else:
                        body(cleaned)
                elif in_keywords and line.strip():
                    keyword_tags(line)
                else:
                    body(line)

        except Exception:
            pass

    path = os.path.join(OUTPUT_DIR, _safe_filename(video_id) + ".pdf")
    pdf.output(path)
    print(f"[Export] PDF saved:      {path}")
    return path


def export_docx(notes_md, video_id):
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        raise RuntimeError("python-docx not installed.")

    doc = Document()
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = Inches(1)
        sec.left_margin = sec.right_margin = Inches(1.2)

    AC  = RGBColor(0x63, 0x66, 0xF1)
    AC2 = RGBColor(0x81, 0x8C, 0xF8)
    BK  = RGBColor(0x1E, 0x1E, 0x1E)

    def add_hr(paragraph):
        p = paragraph._p
        pPr = p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"),   "single")
        bottom.set(qn("w:sz"),    "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "6366F1")
        pBdr.append(bottom)
        pPr.append(pBdr)

    for raw in notes_md.splitlines():
        line = raw.rstrip()
        try:
            if line.startswith("# "):
                p = doc.add_heading(_strip_md(line[2:]), level=1)
                for r in p.runs:
                    r.font.color.rgb = AC
                    r.font.size = Pt(18)
                add_hr(p)

            elif line.startswith("## "):
                p = doc.add_heading(_strip_md(line[3:]), level=2)
                for r in p.runs:
                    r.font.color.rgb = AC
                    r.font.size = Pt(13)

            elif line.startswith("### "):
                p = doc.add_heading(_strip_md(line[4:]), level=3)
                for r in p.runs:
                    r.font.color.rgb = AC2
                    r.font.size = Pt(11)

            elif line.startswith("- "):
                content = re.sub(r"\*\*(.+?)\*\*", r"\1", line[2:])
                p = doc.add_paragraph(style="List Bullet")
                r = p.add_run(content)
                r.font.size      = Pt(11)
                r.font.color.rgb = BK

            elif line.startswith("**") and (":**" in line or ": " in line):
                match = re.match(r"\*\*(.+?)\*\*:?\s*(.*)", line)
                if match:
                    p = doc.add_paragraph()
                    r1 = p.add_run(match.group(1) + ": ")
                    r1.bold = True
                    r1.font.color.rgb = AC2
                    r1.font.size = Pt(10)
                    r2 = p.add_run(match.group(2))
                    r2.font.color.rgb = BK
                    r2.font.size = Pt(10)

            elif line == "---":
                p = doc.add_paragraph()
                add_hr(p)

            elif line == "":
                doc.add_paragraph("")

            else:
                cleaned = re.sub(r"`(.+?)`", r"\1", _strip_md(line))
                if cleaned.strip():
                    p = doc.add_paragraph()
                    r = p.add_run(cleaned)
                    r.font.size      = Pt(10)
                    r.font.color.rgb = BK

        except Exception:
            pass

    path = os.path.join(OUTPUT_DIR, _safe_filename(video_id) + ".docx")
    doc.save(path)
    print(f"[Export] DOCX saved:     {path}")
    return path


def export_all(notes_md, video_id, formats=None):
    if formats is None:
        formats = ["md", "pdf", "docx"]
    paths = {}
    if "md"   in formats: paths["md"]   = export_markdown(notes_md, video_id)
    if "pdf"  in formats: paths["pdf"]  = export_pdf(notes_md, video_id)
    if "docx" in formats: paths["docx"] = export_docx(notes_md, video_id)
    return paths
